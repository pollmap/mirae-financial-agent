#!/usr/bin/env python3
"""Build the repository-owned final master-plan DOCX from the authoritative Markdown."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "18_FINAL_MASTER_PLAN_AND_RELEASE_READINESS.md"
OUTPUT = (
    ROOT
    / "deliverables"
    / "미래에셋_AI_Festival_금융상품_Agent_최종_기획설계개발_마스터플랜.docx"
)

NAVY = "172A46"
TEAL = "007F7B"
SKY = "E8F3F7"
PALE = "F4F7FA"
AMBER = "F4B942"
RED = "B42318"
GRAY = "5D6876"
WHITE = "FFFFFF"


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_text(cell, text: str, *, color: str = NAVY, bold: bool = False, size: float = 8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(_plain(text))
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_margins(cell)


def _plain(value: str) -> str:
    value = value.strip()
    value = value.replace("**", "").replace("`", "")
    value = re.sub(r"\[([^]]+)\]\([^)]+\)", r"\1", value)
    return value


def _set_run_font(run, *, size: float | None = None, color: str | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PAGE ")
    _set_run_font(run, size=8, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def _setup_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(7)

    normal = document.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in (
        ("Title", 30, NAVY),
        ("Subtitle", 13, GRAY),
        ("Heading 1", 22, NAVY),
        ("Heading 2", 15.5, NAVY),
        ("Heading 3", 11.5, TEAL),
    ):
        style = document.styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name.startswith("Heading") or name == "Title"
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(9 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(17.4))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(12.5)
    table.columns[1].width = Cm(4.9)
    _set_cell_text(table.cell(0, 0), "MIRAE ASSET AI FESTIVAL · FINANCIAL PRODUCT AGENT", bold=True, size=7.5)
    _set_cell_text(table.cell(0, 1), "RELEASE READINESS", color=TEAL, bold=True, size=7.5)
    _set_cell_shading(table.cell(0, 0), PALE)
    _set_cell_shading(table.cell(0, 1), SKY)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.add_run("codex/federated-completion-v3  ·  2026-08-09")
    for run in paragraph.runs:
        _set_run_font(run, size=7.5, color=GRAY)
    _page_number(footer.add_paragraph())


def _add_cover(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(42)
    run = paragraph.add_run("FINAL MASTER PLAN")
    _set_run_font(run, size=10, color=TEAL)
    run.bold = True

    title = document.add_paragraph(style="Title")
    title.add_run("금융상품 Agent\n최종 기획·설계·개발 마스터플랜")
    title.paragraph_format.space_after = Pt(14)

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("공식 적합성 · Federated Retrieval · 적대적 검증 · Release Readiness")

    document.add_paragraph("")
    table = document.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(12.7)
    rows = (
        ("현재 판정", "PENDING_EXTERNAL — 로컬 완성 후보, 실서비스 완료 아님"),
        ("기준 브랜치", "codex/federated-completion-v3"),
        ("런타임 소스", "c7c07c9bf6747affd56263e4eb1972e29e72cc56"),
        ("작성일", "2026-08-09 · Asia/Seoul"),
    )
    for row_index, (label, value) in enumerate(rows):
        _set_cell_text(table.cell(row_index, 0), label, color=WHITE, bold=True, size=9)
        _set_cell_shading(table.cell(row_index, 0), NAVY)
        _set_cell_text(table.cell(row_index, 1), value, size=9)
        _set_cell_shading(table.cell(row_index, 1), PALE if row_index % 2 == 0 else WHITE)

    document.add_paragraph("")
    note = document.add_table(rows=1, cols=1)
    note.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = note.cell(0, 0)
    _set_cell_shading(cell, "FFF4E5")
    _set_cell_text(
        cell,
        "정직한 완료선: 실제 HCX·CLOVA Embedding·NCP 공개 HTTPS·사람의 제출 승인 전에는 "
        "‘완벽히 완료’ 또는 ‘실서비스 검증 완료’라고 표시하지 않는다.",
        color=RED,
        bold=True,
        size=9.5,
    )

    document.add_page_break()


def _add_contents(document: Document) -> None:
    document.add_heading("문서 지도", level=1)
    paragraph = document.add_paragraph(
        "이 문서는 표와 실행 증거 중심의 standard business brief 형식으로 구성했다. "
        "수치는 artifacts/release_evidence_v4.json과 일치한다."
    )
    paragraph.paragraph_format.space_after = Pt(10)
    sections = (
        "1. 결론부터",
        "2. 판정 근거의 권위",
        "3. 공식 요구와 실제 구현",
        "4. 최종 아키텍처와 ConditionLedger",
        "5. Federated Retrieval의 정직한 상태",
        "6. 데이터·통화·안전 정책",
        "7. 최종 로컬·fixture 증거",
        "8. 실제 HCX gate",
        "9. Release 상태표",
        "10. 사용자에게만 남는 네 단계",
        "11. 최종 실행 순서",
    )
    for item in sections:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.6)
        paragraph.paragraph_format.first_line_indent = Cm(-0.45)
        paragraph.add_run(item)
    document.add_page_break()


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(rows):
        for column_index in range(width):
            text = row[column_index] if column_index < len(row) else ""
            cell = table.cell(row_index, column_index)
            _set_cell_text(
                cell,
                text,
                color=WHITE if row_index == 0 else NAVY,
                bold=row_index == 0,
                size=8 if width >= 4 else 8.5,
            )
            _set_cell_shading(
                cell,
                NAVY if row_index == 0 else (PALE if row_index % 2 == 0 else WHITE),
            )
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def _add_code(document: Document, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "102033")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(7.8)
        run.font.color.rgb = RGBColor.from_string("EAF2F8")
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def _add_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    first_h2 = True
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("# "):
            index += 1
            continue
        if stripped.startswith("## "):
            if not first_h2:
                document.add_page_break()
            first_h2 = False
            document.add_heading(_plain(stripped[3:]), level=1)
            index += 1
            continue
        if stripped.startswith("### "):
            document.add_heading(_plain(stripped[4:]), level=2)
            index += 1
            continue
        if stripped.startswith("```"):
            code: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index].rstrip())
                index += 1
            _add_code(document, code)
            index += 1
            continue
        if stripped.startswith("|"):
            table_rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    table_rows.append(cells)
                index += 1
            _add_table(document, table_rows)
            continue
        if re.match(r"^[-*] ", stripped):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(_plain(stripped[2:]))
            index += 1
            continue
        if re.match(r"^\d+\. ", stripped):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.6)
            paragraph.paragraph_format.first_line_indent = Cm(-0.45)
            paragraph.add_run(_plain(stripped))
            index += 1
            continue
        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if (
                not candidate
                or candidate.startswith(("#", "|", "```", "- ", "* "))
                or re.match(r"^\d+\. ", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = document.add_paragraph(_plain(" ".join(paragraph_lines)))
        paragraph.paragraph_format.widow_control = True


def main() -> None:
    document = Document()
    _setup_document(document)
    document.core_properties.title = "금융상품 Agent 최종 기획·설계·개발 마스터플랜"
    document.core_properties.subject = "공식 적합성, 설계, 구현, 검증, release readiness"
    document.core_properties.author = "Mirae Financial Agent Team"
    document.core_properties.keywords = "HyperCLOVA X, 금융상품 Agent, Federated Retrieval"
    _add_cover(document)
    _add_contents(document)
    _add_markdown(document, SOURCE.read_text(encoding="utf-8"))
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
